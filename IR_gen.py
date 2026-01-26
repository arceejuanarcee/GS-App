import io
import re
from datetime import date, datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

import ms_graph
import sp_folder_graph as spg


# ==============================
# CONFIG
# ==============================
TEMPLATE_PATH = "Incident Report Template_blank (1).docx"

CITY_CODES = {
    "Davao City": "DVO",
    "Quezon City": "QZN",
}

STANDARD_IMAGE_WIDTH_IN = 5.5

SHAREPOINT_SITE_URL = st.secrets.get("sharepoint", {}).get("site_url", "")
INCIDENT_REPORTS_ROOT_PATH = st.secrets.get("sharepoint", {}).get(
    "incident_reports_root_path",
    "Ground Station Operations/Installations, Maintenance and Repair/Incident Reports",
)

SCOPES = ms_graph.DEFAULT_SCOPES_WRITE


# ==============================
# DOCX HELPERS
# ==============================
def _clear_table_rows_except_header(table, header_rows=1):
    while len(table.rows) > header_rows:
        tbl = table._tbl
        tr = table.rows[-1]._tr
        tbl.remove(tr)


def _set_2col_table_value(table, label, value):
    for row in table.rows:
        if row.cells[0].text.strip() == label.strip():
            row.cells[1].text = "" if value is None else str(value)
            return


def _set_paragraph_after_heading(doc, heading_text, new_text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text.strip():
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].text = new_text or ""
            return


def _insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _append_figures_after_heading(doc, heading_text, files, captions, figure_start, section_label):
    if not files:
        return figure_start

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text.strip():
            anchor = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else p
            fig_no = figure_start

            for idx, f in enumerate(files):
                img_p = _insert_paragraph_after(anchor)
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(io.BytesIO(f.getvalue()), width=Inches(STANDARD_IMAGE_WIDTH_IN))

                caption_text = ""
                if captions and idx < len(captions):
                    caption_text = (captions[idx] or "").strip()
                if not caption_text:
                    caption_text = f.name.rsplit(".", 1)[0]

                cap_p = _insert_paragraph_after(img_p)
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_p.add_run(f"Figure {fig_no}. {section_label} – {caption_text}")
                cap_run.italic = True

                anchor = cap_p
                fig_no += 1

            return fig_no

    return figure_start


def _fill_sequence_table(table, df):
    _clear_table_rows_except_header(table, header_rows=0)
    for _, r in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r.get("Date", ""))
        cells[1].text = str(r.get("Time", ""))
        cells[2].text = str(r.get("Category", ""))
        cells[3].text = str(r.get("Message", ""))


def _fill_actions_table(table, df):
    _clear_table_rows_except_header(table, header_rows=1)
    for _, r in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r.get("Date", ""))
        cells[1].text = str(r.get("Time", ""))
        cells[2].text = str(r.get("Performed by", ""))
        cells[3].text = str(r.get("Action", ""))
        cells[4].text = str(r.get("Result", ""))


def generate_docx(data):
    doc = Document(TEMPLATE_PATH)
    t0, t1, t2, t3 = doc.tables[0], doc.tables[1], doc.tables[2], doc.tables[3]

    _set_2col_table_value(t0, "Reported by", data["reported_by"])
    _set_2col_table_value(t0, "Position", data["position"])
    _set_2col_table_value(t0, "Date of Report", data["date_of_report"])
    _set_2col_table_value(t0, "Incident No.", data["full_incident_no"])

    _set_2col_table_value(t1, "Date (YYYY-MM-DD)", data["incident_date"])
    _set_2col_table_value(t1, "Time", data["incident_time"])
    _set_2col_table_value(t1, "Location", data["location"])
    _set_2col_table_value(t1, "Current Status", data["current_status"])

    _set_paragraph_after_heading(doc, "Nature of Incident", data["nature"])
    _set_paragraph_after_heading(doc, "Damages Incurred (if any)", data["damages"])
    _set_paragraph_after_heading(doc, "Investigation and Analysis", data["investigation"])
    _set_paragraph_after_heading(doc, "Conclusion and Recommendations", data["conclusion"])

    _fill_sequence_table(t2, data["sequence_df"])
    _fill_actions_table(t3, data["actions_df"])

    fig = 1
    fig = _append_figures_after_heading(doc, "Sequence of Events", data["sequence_images"], data["sequence_captions"], fig, "Sequence of Events")
    fig = _append_figures_after_heading(doc, "Damages Incurred (if any)", data["damages_images"], data["damages_captions"], fig, "Damages Incurred")
    fig = _append_figures_after_heading(doc, "Investigation and Analysis", data["investigation_images"], data["investigation_captions"], fig, "Investigation and Analysis")
    fig = _append_figures_after_heading(doc, "Conclusion and Recommendations", data["conclusion_images"], data["conclusion_captions"], fig, "Conclusion and Recommendations")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# ==============================
# UI HELPERS
# ==============================
def captions_editor(files, key):
    if not files:
        return []
    df = pd.DataFrame({"File": [f.name for f in files], "Caption": ["" for _ in files]})
    edited = st.data_editor(df, key=key, num_rows="fixed", use_container_width=True)
    return edited["Caption"].tolist()


def normalize_serial(serial_raw: str) -> str:
    s = (serial_raw or "").strip()
    if not s:
        return ""
    if not re.fullmatch(r"\d{1,4}", s):
        return ""
    return s.zfill(4)


def _must_have_token():
    ms_graph.login_ui(scopes=SCOPES)
    token = ms_graph.get_access_token()
    if not token:
        st.stop()
    return token


# ==============================
# STREAMLIT APP
# ==============================
st.set_page_config(page_title="Incident Report Generator", layout="wide")
st.title("Incident Report Generator")

token = _must_have_token()

if not SHAREPOINT_SITE_URL:
    st.error("Missing sharepoint.site_url in Streamlit secrets.")
    st.stop()

if "sp_site_id" not in st.session_state or "sp_drive_id" not in st.session_state:
    with st.spinner("Resolving SharePoint site/drive..."):
        st.session_state["sp_site_id"] = spg.resolve_site_id(token, SHAREPOINT_SITE_URL)
        st.session_state["sp_drive_id"] = spg.get_default_drive_id(token, st.session_state["sp_site_id"])

drive_id = st.session_state["sp_drive_id"]

this_year = str(datetime.now().year)
year = st.selectbox("Year folder", [this_year, str(int(this_year) - 1)], index=0)

city = st.selectbox("Ground Station Location", list(CITY_CODES.keys()))
site_code = CITY_CODES[city]

serial_raw = st.text_input("Incident serial (000#)", placeholder="e.g. 0001 (or 1)")
serial = normalize_serial(serial_raw)
full_incident_no = f"SMCOD-IR-GS-{site_code}-{year}-{serial}" if serial else ""
st.text_input("Full Incident No. (auto)", value=full_incident_no, disabled=True)

with st.form("ir_form"):
    c1, c2 = st.columns(2)
    with c1:
        reported_by = st.text_input("Reported by")
        position = st.text_input("Position")
        date_of_report = st.date_input("Date of Report", value=date.today()).strftime("%Y-%m-%d")
    with c2:
        incident_date = st.date_input("Incident Date", value=date.today()).strftime("%Y-%m-%d")
        incident_time = st.text_input("Incident Time", value=datetime.now().strftime("%H:%M:%S"))
        location = st.text_input("Location", value=city)
        current_status = st.selectbox("Current Status", ["Resolved", "Ongoing", "Monitoring", "Open"])

    nature = st.text_area("Nature of Incident", height=120)

    st.subheader("Sequence of Events")
    seq_df = st.data_editor(
        pd.DataFrame([{"Date": incident_date, "Time": "", "Category": "", "Message": ""}]),
        num_rows="dynamic",
        use_container_width=True,
    )
    seq_imgs = st.file_uploader("Sequence Photos", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    seq_caps = captions_editor(seq_imgs or [], "seq_caps")

    damages = st.text_area("Damages Incurred", value="None")
    dmg_imgs = st.file_uploader("Damage Photos", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    dmg_caps = captions_editor(dmg_imgs or [], "dmg_caps")

    investigation = st.text_area("Investigation and Analysis", height=120)
    inv_imgs = st.file_uploader("Investigation Photos", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    inv_caps = captions_editor(inv_imgs or [], "inv_caps")

    conclusion = st.text_area("Conclusion and Recommendations", height=120)
    con_imgs = st.file_uploader("Conclusion Photos", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    con_caps = captions_editor(con_imgs or [], "con_caps")

    st.subheader("Response and Actions Taken")
    actions_df = st.data_editor(
        pd.DataFrame([{"Date": incident_date, "Time": "", "Performed by": "", "Action": "", "Result": ""}]),
        num_rows="dynamic",
        use_container_width=True,
    )

    submit = st.form_submit_button("Generate Report")

if submit:
    if not serial:
        st.error("Enter a valid incident serial (numbers only up to 4 digits). Example: 0001 or 1.")
        st.stop()

    try:
        is_dup = spg.check_duplicate_ir(token, drive_id, INCIDENT_REPORTS_ROOT_PATH, year, city, full_incident_no)
        if is_dup:
            st.error("Duplicate found: this Incident No folder already exists. Use a new serial.")
            st.stop()
    except Exception as e:
        st.warning(f"Duplicate check failed (continuing): {e}")

    data = {
        "reported_by": reported_by,
        "position": position,
        "date_of_report": date_of_report,
        "full_incident_no": full_incident_no,
        "incident_date": incident_date,
        "incident_time": incident_time,
        "location": location,
        "current_status": current_status,
        "nature": nature,
        "damages": damages,
        "investigation": investigation,
        "conclusion": conclusion,
        "sequence_df": seq_df,
        "actions_df": actions_df,
        "sequence_images": seq_imgs or [],
        "damages_images": dmg_imgs or [],
        "investigation_images": inv_imgs or [],
        "conclusion_images": con_imgs or [],
        "sequence_captions": seq_caps or [],
        "damages_captions": dmg_caps or [],
        "investigation_captions": inv_caps or [],
        "conclusion_captions": con_caps or [],
    }

    docx_bytes = generate_docx(data)

    try:
        with st.spinner("Creating folder and uploading DOCX to SharePoint..."):
            incident_folder = spg.ensure_path(
                token,
                drive_id,
                INCIDENT_REPORTS_ROOT_PATH,
                parts=[year, city, full_incident_no],
            )

            spg.upload_file_to_folder(
                token,
                drive_id,
                folder_item_id=incident_folder["id"],
                filename=f"{full_incident_no}.docx",
                content_bytes=docx_bytes,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        st.success("Report generated, folder created, and DOCX uploaded.")
    except Exception as e:
        st.error(f"Upload failed: {e}")

    st.download_button(
        "Download Incident Report (DOCX)",
        data=docx_bytes,
        file_name=f"{full_incident_no}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ==============================
# FILE VIEWER & EDITOR (FIXED FOR YOUR STRUCTURE)
# ==============================
st.divider()
st.subheader("File Viewer & Editor (SharePoint)")

viewer_year = st.selectbox("Year", [this_year, str(int(this_year) - 1)], index=0, key="v_year")
viewer_city = st.selectbox("Ground Station Location", list(CITY_CODES.keys()), key="v_city")

base_path = f"{INCIDENT_REPORTS_ROOT_PATH}/{viewer_year}/{viewer_city}"
st.caption(f"City folder: {base_path}")

# Load incident folders
if st.button("Refresh folders/files", key="v_refresh"):
    for k in ["v_incident_folders", "v_files", "v_loaded_id", "v_loaded_text"]:
        st.session_state.pop(k, None)

if "v_incident_folders" not in st.session_state:
    try:
        st.session_state["v_incident_folders"] = spg.list_incident_folders(token, drive_id, base_path)
    except Exception as e:
        st.error(f"Cannot list incident folders: {e}")
        st.session_state["v_incident_folders"] = []

folders = st.session_state.get("v_incident_folders", [])
folder_names = [f["name"] for f in folders]

incident_folder_name = st.selectbox("Incident Folder (Incident No.)", ["-- select --"] + folder_names, key="v_folder")

if incident_folder_name != "-- select --":
    folder_meta = next((x for x in folders if x["name"] == incident_folder_name), None)
    folder_id = folder_meta["id"]

    if "v_files" not in st.session_state or st.session_state.get("v_files_folder") != folder_id:
        try:
            st.session_state["v_files"] = spg.list_files(token, drive_id, folder_id)
            st.session_state["v_files_folder"] = folder_id
        except Exception as e:
            st.error(f"Cannot list files in incident folder: {e}")
            st.session_state["v_files"] = []

    files = st.session_state.get("v_files", [])
    file_names = [x["name"] for x in files]

    selected_file = st.selectbox("File", ["-- select --"] + file_names, key="v_file")

    if selected_file != "-- select --":
        fmeta = next((x for x in files if x["name"] == selected_file), None)
        file_id = fmeta["id"]
        ext = selected_file.lower().rsplit(".", 1)[-1] if "." in selected_file else ""

        # DOCX / others: download
        if st.button("Download file", key="v_download"):
            b = spg.download_file_bytes(token, drive_id, file_id)
            st.download_button(
                "Click to download",
                data=b,
                file_name=selected_file,
                mime="application/octet-stream",
            )

        # Text files: view + edit + save
        if ext in ["txt", "md", "log", "csv"]:
            if st.button("Load contents", key="v_load"):
                try:
                    st.session_state["v_loaded_id"] = file_id
                    st.session_state["v_loaded_text"] = spg.download_file_text(token, drive_id, file_id)
                except Exception as e:
                    st.error(f"Load failed: {e}")

            if st.session_state.get("v_loaded_id") == file_id:
                new_text = st.text_area("Contents", value=st.session_state.get("v_loaded_text", ""), height=320)

                c1, c2 = st.columns([1, 1])
                with c1:
                    if st.button("Save (overwrite)", key="v_save"):
                        try:
                            spg.update_file_text(token, drive_id, file_id, new_text)
                            st.success("Saved.")
                            st.session_state["v_loaded_text"] = new_text
                        except Exception as e:
                            st.error(f"Save failed: {e}")

                with c2:
                    if st.button("Unload", key="v_unload"):
                        st.session_state.pop("v_loaded_id", None)
                        st.session_state.pop("v_loaded_text", None)
        else:
            st.info("Editing is enabled only for text files (.txt, .md, .log, .csv). For DOCX/PDF, use Download.")
